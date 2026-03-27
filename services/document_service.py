# services/document_service.py — PLAGENOR 4.0 Document Generation
# Generates: IBTIKAR request forms, Platform notes, GENOCLAB invoices/quotes

from __future__ import annotations
import os, shutil
from datetime import datetime
from typing import Optional
import config


def _get_template_dir():
    return os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "templates")


def _get_output_dir():
    d = os.path.join(config.DATA_DIR, "reports")
    os.makedirs(d, exist_ok=True)
    return d


# ═══════════════════════════════════════════════════════════════════════════
# IBTIKAR REQUEST FORM (DOCX) — Pre-filled from request data
# ═══════════════════════════════════════════════════════════════════════════
def generate_ibtikar_form(request: dict, service: dict = None, requester: dict = None) -> str:
    """
    Generate pre-filled IBTIKAR form by CLONING the official template DOCX
    and replacing placeholder text with submitted data.
    Preserves all original formatting, headers, logos, and structure.
    """
    try:
        from docx import Document
        import shutil

        svc_code = (service or {}).get("code", "") or (service or {}).get("service_code", "") or request.get("service_code", "")

        # Map service codes to template files
        template_map = {
            "EGTP-IMT": "egtp_imt.docx",
            "EGTP-Seq01": "egtp_seqs.docx",
            "EGTP-Seq02": "egtp_seq02.docx",
            "EGTP-SeqS": "egtp_seqs.docx",
            "EGTP-PCR": "egtp_pcr.docx",
            "EGTP-CAN": "egtp_can.docx",
            "EGTP-GDE": "egtp_gde.docx",
            "EGTP-PS": "egtp_ps.docx",
            "EGTP-Lyoph": "egtp_lyoph.docx",
            "EGTP-Illumina-Microbial-WGS": "egtp_illumina_microbial-wgs.docx",
        }

        template_file = template_map.get(svc_code, "")
        template_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "templates", "ibtikar")
        template_path = os.path.join(template_dir, template_file)

        if not template_file or not os.path.exists(template_path):
            # Fallback: generate basic form if no template found
            return _generate_basic_ibtikar_form(request, service, requester)

        # ── CLONE the official template ──────────────────────────────────
        req_id = request.get("id", "")[:8].upper()
        year = datetime.now().year
        filename = f"IBTIKAR_{svc_code}_{req_id}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = os.path.join(_get_output_dir(), filename)
        shutil.copy2(template_path, output_path)

        # ── OPEN and REPLACE placeholders ────────────────────────────────
        doc = Document(output_path)

        # Get data from request
        req_info = request.get("requester", {})
        svc_params = request.get("service_params", {})
        if not req_info and requester:
            req_info = {"full_name": requester.get("full_name", ""), "email": requester.get("email", "")}

        # Build replacement map for BOTH placeholder styles
        replacements = {
            # {{TAG}} style (EGTP-IMT)
            "{{FULL_NAME}}": req_info.get("full_name", ""),
            "{{ETABLISSEMENT}}": req_info.get("institution", ""),
            "{{LABORATORY}}": req_info.get("laboratory", ""),
            "{{FUNCTION}}": req_info.get("status", ""),
            "{{PHONE}}": req_info.get("phone", ""),
            "{{EMAIL}}": req_info.get("email", ""),
            "{{ANALYSIS_FRAME}}": svc_params.get("analysis_frame", ""),
            "{{PROJECT_TITLE}}": svc_params.get("project_title", request.get("title", "")),
            "{{PROJECT_DIRECTOR}}": req_info.get("supervisor", ""),
            "{{ANALYSIS_MODE_LABEL}}": svc_params.get("analysis_mode", ""),
            "{{MALDI_TARGET_TYPE}}": svc_params.get("maldi_target_type", ""),
            "{{FRESH_CULTURE}}": "Oui" if svc_params.get("fresh_culture_available") else "Non",
            # Descriptive text style (all other forms)
            "Nom complet du demandeur": req_info.get("full_name", ""),
            "Nom de l\u2019établissement": req_info.get("institution", ""),
            "Nom de l\'établissement": req_info.get("institution", ""),
            "Nom du laboratoire": req_info.get("laboratory", ""),
            "Etudiant/doctorant/Chercheur/MCA\u2026": req_info.get("status", ""),
            "Etudiant/doctorant/Chercheur/MCA…": req_info.get("status", ""),
            "E-mail du demandeur": req_info.get("email", ""),
            "Numéro du demandeur": req_info.get("phone", ""),
            "Choisissez un élément.": svc_params.get("analysis_frame", ""),
            "Titre de la recherche": svc_params.get("project_title", request.get("title", "")),
            "Nom et prénom du directeur de recherche encadrant le projet": req_info.get("supervisor", ""),
            "Nom et prénom du directeur de recherche\nencadrant le projet": req_info.get("supervisor", ""),
            # Request number
            "……./2026": f"{req_id}/{year}",
            "\u2026\u2026./2026": f"{req_id}/{year}",
        }

        # Replace in all paragraphs
        for paragraph in doc.paragraphs:
            _replace_in_paragraph(paragraph, replacements)

        # Handle label-only fields (no placeholder, just "Label : *")
        # These need the value APPENDED to the last run
        append_map = {
            "Fonction / Poste": req_info.get("status", ""),
        }
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            for label, value in append_map.items():
                if label in text and value and "{{" not in text:
                    # Check if this paragraph only has the label (no data yet)
                    if text.endswith("*") or text.endswith("* "):
                        if paragraph.runs:
                            paragraph.runs[-1].text = paragraph.runs[-1].text.rstrip() + " " + str(value)

        # Replace in all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, replacements)

        # ── FILL sample table rows if data available ─────────────────────
        sample_table_data = request.get("sample_table", [])
        if sample_table_data and doc.tables:
            _fill_sample_table(doc, sample_table_data)

        doc.save(output_path)

        # Record document
        from core.repository import save_document
        save_document({
            "type": "IBTIKAR_FORM",
            "request_id": request.get("id", ""),
            "service_code": svc_code,
            "filename": filename,
            "filepath": output_path,
            "template_used": template_file,
            "created_at": datetime.now().isoformat(),
        })

        return output_path
    except Exception as e:
        print(f"IBTIKAR form generation error: {e}")
        import traceback; traceback.print_exc()
        return ""


def _generate_basic_ibtikar_form(request, service, requester):
    """Fallback: generate a basic form when no official template is available."""
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading(f"IBTIKAR — {(service or {}).get('service_name', '')}", level=1)
        doc.add_paragraph(f"Réf: {request.get('id','')[:8]}/{datetime.now().year}/IBTIKAR")
        doc.add_paragraph(f"Demandeur: {request.get('requester_name','')}")
        doc.add_paragraph(f"Service: {(service or {}).get('code','')}")
        filename = f"IBTIKAR_BASIC_{request.get('id','')[:8]}.docx"
        path = os.path.join(_get_output_dir(), filename)
        doc.save(path)
        return path
    except Exception:
        return ""


def _replace_in_paragraph(paragraph, replacements: dict):
    """Replace placeholder text in a paragraph while preserving formatting."""
    full_text = paragraph.text
    if not full_text:
        return

    for old_text, new_text in replacements.items():
        if old_text in full_text:
            # Try to replace in individual runs first (preserves formatting)
            replaced = False
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, str(new_text) if new_text else "")
                    replaced = True

            # If not found in individual runs, the text spans multiple runs
            # Rebuild paragraph text
            if not replaced and old_text in full_text:
                # Collect all runs, replace in combined text, redistribute
                combined = "".join(run.text for run in paragraph.runs)
                if old_text in combined:
                    new_combined = combined.replace(old_text, str(new_text) if new_text else "")
                    # Put all text in first run, clear others
                    if paragraph.runs:
                        paragraph.runs[0].text = new_combined
                        for run in paragraph.runs[1:]:
                            run.text = ""


def _fill_sample_table(doc, samples: list):
    """Fill the sample table in the document with submitted sample data."""
    # Find the largest table (the sample table)
    sample_table = None
    max_cols = 0
    for table in doc.tables:
        if len(table.columns) > max_cols:
            max_cols = len(table.columns)
            sample_table = table

    if not sample_table or max_cols < 3:
        return

    # Define column mapping: ordered list of sample dict keys matching table columns
    # Column 0 = row number (N), then data columns
    column_keys = [
        "sample_code", "organism_type", "sample_origin", "isolation_date",
        "culture_medium", "culture_conditions", "remarks",
        # Extra keys for other services
        "target_gene", "dna_origin", "dna_type", "primer_name",
        "primer_sequence", "primer_length",
    ]

    # Fill data rows (skip header row)
    data_rows = sample_table.rows[1:]
    for idx, sample in enumerate(samples):
        if idx >= len(data_rows):
            break

        row = data_rows[idx]
        cells = row.cells

        # Column 0 = row number
        if len(cells) > 0:
            cells[0].paragraphs[0].text = str(idx + 1).zfill(2)

        # Fill remaining columns with sample data
        col_idx = 1
        for key in column_keys:
            if col_idx >= len(cells):
                break
            val = sample.get(key)
            if val is not None and str(val).strip():
                # Clear existing text and set new value
                for p in cells[col_idx].paragraphs:
                    if not p.text.strip():  # Only fill empty cells
                        if p.runs:
                            p.runs[0].text = str(val)
                        else:
                            p.text = str(val)
                        break
                col_idx += 1
            elif key in sample:
                col_idx += 1  # Skip empty optional fields but advance column


def generate_platform_note(request: dict, actor: dict, service: dict = None) -> str:
    """Generate Platform Note DOCX for admin validation."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        req_id = request.get("id", "")[:8].upper()
        year = datetime.now().year
        svc_name = (service or {}).get("name", request.get("service_id", ""))
        svc_code = (service or {}).get("code", "")
        channel = request.get("channel", "IBTIKAR")

        # --- Header ---
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_p.add_run("RÉPUBLIQUE ALGÉRIENNE DÉMOCRATIQUE ET POPULAIRE")
        run.font.size = Pt(10)
        run.font.bold = True

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run("MINISTÈRE DE L'ENSEIGNEMENT SUPÉRIEUR ET DE LA RECHERCHE SCIENTIFIQUE")
        run.font.size = Pt(9)

        inst = doc.add_paragraph()
        inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = inst.add_run(config.PLATFORM_INSTITUTION)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

        doc.add_paragraph()

        # --- Title ---
        title = doc.add_heading("NOTE DE PLATEFORME", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

        # Reference
        ref_p = doc.add_paragraph()
        ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ref_p.add_run(f"Réf: NOTE-PLT-{req_id}/{year}/{channel}")
        run.font.size = Pt(11)
        run.font.bold = True

        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
        run.font.size = Pt(10)

        doc.add_paragraph()

        # --- Object ---
        obj = doc.add_paragraph()
        obj.add_run("Objet: ").bold = True
        obj.add_run(f"Validation de la demande d'analyse — {svc_name}")

        doc.add_paragraph()

        # --- Request Details ---
        doc.add_heading("Détails de la Demande", level=2)

        details = [
            ("N° Demande", f"{req_id}/{year}/{channel}"),
            ("Canal", channel),
            ("Service demandé", f"{svc_name} ({svc_code})"),
            ("Titre du projet", request.get("title", "")),
            ("Demandeur", request.get("requester_name", "")),
            ("Nombre d'échantillons", str(request.get("sample_count", ""))),
            ("Date de soumission", request.get("created_at", "")[:10]),
        ]

        if channel == "IBTIKAR" and request.get("budget_amount"):
            details.append(("Budget demandé", f"{request.get('budget_amount', 0):,.0f} DZD"))

        table = doc.add_table(rows=len(details), cols=2)
        table.style = 'Light Grid Accent 1'
        for i, (label, value) in enumerate(details):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            for p in table.rows[i].cells[0].paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        doc.add_paragraph()

        # --- Description ---
        doc.add_heading("Description de la Demande", level=2)
        doc.add_paragraph(request.get("description", "Aucune description fournie."))

        doc.add_paragraph()

        # --- Validation ---
        doc.add_heading("Avis de la Plateforme", level=2)

        val_table = doc.add_table(rows=4, cols=2)
        val_table.style = 'Light Grid Accent 1'
        val_fields = [
            ("Validé par", actor.get("full_name", actor.get("username", ""))),
            ("Rôle", config.ROLE_LABELS.get(actor.get("role", ""), "")),
            ("Date de validation", datetime.now().strftime("%d/%m/%Y %H:%M")),
            ("Décision", "✅ VALIDÉE — La demande est conforme et peut être traitée."),
        ]
        for i, (label, value) in enumerate(val_fields):
            val_table.rows[i].cells[0].text = label
            val_table.rows[i].cells[1].text = str(value)

        doc.add_paragraph()

        # --- Budget section for IBTIKAR ---
        if channel == "IBTIKAR" and request.get("budget_amount"):
            doc.add_heading("Volet Financier (IBTIKAR)", level=2)
            from core.financial_engine import get_ibtikar_budget_used
            used = get_ibtikar_budget_used()
            cap = config.IBTIKAR_BUDGET_CAP
            amount = request.get("budget_amount", 0)

            budget_table = doc.add_table(rows=4, cols=2)
            budget_table.style = 'Light Grid Accent 1'
            budget_fields = [
                ("Plafond annuel", f"{cap:,.0f} DZD"),
                ("Budget consommé", f"{used:,.0f} DZD"),
                ("Montant demandé", f"{amount:,.0f} DZD"),
                ("Budget restant après approbation", f"{max(0, cap - used - amount):,.0f} DZD"),
            ]
            for i, (label, value) in enumerate(budget_fields):
                budget_table.rows[i].cells[0].text = label
                budget_table.rows[i].cells[1].text = value

        doc.add_paragraph()

        # --- Signature ---
        doc.add_heading("Signature", level=2)

        sig = doc.add_paragraph()
        sig.add_run(f"\n\n{actor.get('full_name', '')}\n")
        sig.add_run(f"{config.ROLE_LABELS.get(actor.get('role', ''), '')}\n")
        sig.add_run(f"{config.PLATFORM_INSTITUTION}\n")

        # --- Footer ---
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"Document généré automatiquement par PLAGENOR {config.PLATFORM_VERSION}\n"
            f"© {config.PLATFORM_YEAR} {config.PLATFORM_AUTHOR}\n"
            f"{config.PLATFORM_ADDRESS} · Tél: {config.PLATFORM_PHONE}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)

        # Save
        filename = f"NOTE_PLT_{req_id}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        filepath = os.path.join(_get_output_dir(), filename)
        doc.save(filepath)

        from core.repository import save_document
        save_document({
            "type": "PLATFORM_NOTE",
            "request_id": request.get("id", ""),
            "channel": channel,
            "filename": filename,
            "filepath": filepath,
            "created_by": actor.get("id", ""),
            "created_at": datetime.now().isoformat(),
        })

        return filepath
    except Exception as e:
        print(f"Platform note generation error: {e}")
        return ""


# ═══════════════════════════════════════════════════════════════════════════
# GENOCLAB QUOTE (DEVIS) — Pre-filled from request + service pricing
# ═══════════════════════════════════════════════════════════════════════════
def generate_genoclab_quote(request: dict, service: dict, actor: dict) -> str:
    """Generate GENOCLAB quote DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        req_id = request.get("id", "")[:8].upper()
        year = datetime.now().year
        svc_name = service.get("name", "")
        svc_code = service.get("code", "")

        # Header
        h = doc.add_heading("DEVIS — GENOCLAB", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x11, 0x7A, 0x65)

        inst = doc.add_paragraph()
        inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = inst.add_run(config.PLATFORM_INSTITUTION)
        run.font.size = Pt(10)

        doc.add_paragraph()

        ref = doc.add_paragraph()
        ref.add_run(f"Réf: DEVIS-{req_id}/{year}/GENOCLAB").bold = True

        date_p = doc.add_paragraph()
        date_p.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}")

        doc.add_paragraph()

        # Client info
        doc.add_heading("Client", level=2)
        client_info = [
            ("Nom/Organisation", request.get("client_name", request.get("organization", ""))),
            ("Contact", request.get("contact", "")),
        ]
        for label, val in client_info:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(val))

        doc.add_paragraph()

        # Service details
        doc.add_heading("Prestation", level=2)

        samples = request.get("sample_count", 1)
        price = service.get("base_price", service.get("price", 0))
        subtotal = price * samples
        vat = round(subtotal * config.VAT_RATE, 2)
        total = round(subtotal + vat, 2)

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        rows_data = [
            ("Service", f"{svc_name} ({svc_code})"),
            ("Nombre d'échantillons", str(samples)),
            ("Prix unitaire", f"{price:,.0f} DZD"),
            ("Sous-total HT", f"{subtotal:,.0f} DZD"),
            (f"TVA ({config.VAT_RATE*100:.0f}%)", f"{vat:,.0f} DZD"),
        ]
        for i, (label, val) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = val

        doc.add_paragraph()

        total_p = doc.add_paragraph()
        total_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = total_p.add_run(f"TOTAL TTC: {total:,.0f} DZD")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x11, 0x7A, 0x65)

        doc.add_paragraph()

        # Terms
        doc.add_heading("Conditions", level=2)
        terms = [
            "Ce devis est valable 30 jours à compter de sa date d'émission.",
            "Le paiement est exigé avant le début de l'analyse.",
            "Les résultats seront transmis dans le délai indiqué après réception du paiement.",
        ]
        for t in terms:
            doc.add_paragraph(f"• {t}")

        doc.add_paragraph()

        # Signatures
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.rows[0].cells[0].text = "Pour GENOCLAB — ESSBO"
        sig_table.rows[0].cells[1].text = "Acceptation du client"
        sig_table.rows[1].cells[0].text = f"\n\n{actor.get('full_name', '')}\nDate: {datetime.now().strftime('%d/%m/%Y')}"
        sig_table.rows[1].cells[1].text = "\n\nSignature:\nDate:"

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"GENOCLAB · {config.PLATFORM_INSTITUTION} · {config.PLATFORM_ADDRESS}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)

        # Save
        filename = f"DEVIS_GENOCLAB_{req_id}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(_get_output_dir(), filename)
        doc.save(filepath)

        from core.repository import save_document
        save_document({
            "type": "GENOCLAB_QUOTE",
            "request_id": request.get("id", ""),
            "filename": filename,
            "filepath": filepath,
            "created_by": actor.get("id", ""),
            "created_at": datetime.now().isoformat(),
        })

        # Store quote amount on request
        request["quote_amount"] = total

        return filepath
    except Exception as e:
        print(f"Quote generation error: {e}")
        return ""


# ═══════════════════════════════════════════════════════════════════════════
# GENOCLAB INVOICE (FACTURE)
# ═══════════════════════════════════════════════════════════════════════════
def generate_invoice_document(invoice: dict, request: dict, client_name: str = "") -> str:
    """Generate invoice DOCX for GENOCLAB."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        h = doc.add_heading("FACTURE", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x11, 0x7A, 0x65)

        inst = doc.add_paragraph()
        inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = inst.add_run(f"GENOCLAB · {config.PLATFORM_INSTITUTION}")
        run.font.size = Pt(10)

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(f"N° Facture: {invoice.get('invoice_number', '')}").bold = True

        p2 = doc.add_paragraph()
        p2.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}")

        p3 = doc.add_paragraph()
        p3.add_run(f"Client: {client_name or request.get('client_name', '')}").bold = True

        doc.add_paragraph()
        doc.add_heading("Détails", level=2)

        items = invoice.get("line_items", [])
        if items:
            table = doc.add_table(rows=len(items)+1, cols=4)
            table.style = 'Light Grid Accent 1'
            for i, hdr in enumerate(["Description", "Qté", "Prix unitaire", "Total"]):
                table.rows[0].cells[i].text = hdr
            for idx, item in enumerate(items):
                qty = item.get("quantity", 1)
                price = item.get("unit_price", 0)
                table.rows[idx+1].cells[0].text = item.get("description", "")
                table.rows[idx+1].cells[1].text = str(qty)
                table.rows[idx+1].cells[2].text = f"{price:,.0f} DZD"
                table.rows[idx+1].cells[3].text = f"{qty*price:,.0f} DZD"

        doc.add_paragraph()

        totals = doc.add_paragraph()
        totals.add_run(f"Sous-total HT: {invoice.get('subtotal_ht',0):,.0f} DZD\n")
        totals.add_run(f"TVA ({invoice.get('vat_rate',0.19)*100:.0f}%): {invoice.get('vat_amount',0):,.0f} DZD\n")
        run = totals.add_run(f"TOTAL TTC: {invoice.get('total_ttc',0):,.0f} DZD")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x11, 0x7A, 0x65)

        doc.add_paragraph()

        lock = doc.add_paragraph()
        lock.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = lock.add_run("🔒 Ce document est verrouillé et ne peut être modifié.")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"{config.PLATFORM_INSTITUTION}\n{config.PLATFORM_ADDRESS}\nTél: {config.PLATFORM_PHONE}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)

        inv_num = invoice.get("invoice_number", "INV").replace("/", "-")
        filename = f"{inv_num}.docx"
        filepath = os.path.join(config.INVOICES_DIR, filename)
        doc.save(filepath)

        from core.repository import save_document
        save_document({
            "type": "INVOICE",
            "invoice_id": invoice.get("id", ""),
            "request_id": request.get("id", ""),
            "filename": filename,
            "filepath": filepath,
            "created_at": datetime.now().isoformat(),
        })
        return filepath
    except Exception as e:
        print(f"Invoice document error: {e}")
        return ""
